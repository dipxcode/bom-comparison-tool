import re
import io
import logging
import pandas as pd
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

KEY_COLUMN_ALIASES = [
    'part_number','part_no','partno','part#','pn','p/n','p_n',
    'item_number','item_no','itemno','item#','item',
    'mpn','mfr_pn','mfr_part_no','manufacturer_part_number',
    'component','component_id','comp_id',
    'reference','ref_des','refdes','reference_designator',
    'sku','material','material_number','mat_no',
]


def normalize_col(name: str) -> str:
    return re.sub(r'[\s\-/]+', '_', str(name).strip().lower())


def read_bom_file(file_path: str) -> pd.DataFrame:
    path = Path(str(file_path))
    ext  = path.suffix.lower()
    if ext in ('.xlsx', '.xls'):
        engine = 'openpyxl' if ext == '.xlsx' else 'xlrd'
        df = pd.read_excel(path, dtype=str, keep_default_na=False, engine=engine)
    elif ext == '.csv':
        df = _read_csv(path)
    elif ext == '.txt':
        try:
            df = pd.read_csv(path, sep='\t', dtype=str, keep_default_na=False)
            if df.shape[1] < 2:
                raise ValueError("Not tab-separated")
        except Exception:
            df = _read_csv(path)
    elif ext == '.docx':
        df = _read_docx(path)
    elif ext == '.pdf':
        df = _read_pdf(path)
    else:
        raise ValueError(f'Unsupported file format: {ext}')

    df.dropna(how='all', inplace=True)
    df.dropna(axis=1, how='all', inplace=True)
    df.columns = [str(c).strip() for c in df.columns]
    return df


def _read_csv(path: Path) -> pd.DataFrame:
    for enc in ('utf-8', 'latin-1', 'cp1252'):
        try:
            return pd.read_csv(path, encoding=enc, dtype=str, keep_default_na=False)
        except UnicodeDecodeError:
            continue
    raise ValueError(f'Cannot decode CSV: {path}')


def _read_docx(path: Path) -> pd.DataFrame:
    from docx import Document
    doc = Document(str(path))
    if not doc.tables:
        raise ValueError('No tables found in DOCX file.')
    table = doc.tables[0]
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not rows:
        raise ValueError('DOCX table is empty.')
    headers = rows[0]
    data    = rows[1:]
    return pd.DataFrame(data, columns=headers, dtype=str)


def _read_pdf(path: Path) -> pd.DataFrame:
    import pdfplumber

    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            try:
                tables = page.extract_tables()
            except Exception:
                tables = []
            if tables:
                for tbl in tables:
                    if not tbl or len(tbl) < 2:
                        continue
                    raw_headers = tbl[0]
                    if not any(raw_headers):
                        continue
                    headers = [str(h).strip() if h else f'Col{i}' for i, h in enumerate(raw_headers)]
                    data = [ [str(c).strip() if c else '' for c in row] for row in tbl[1:] if any(c for c in row if c) ]

                    if data:
                        df = pd.DataFrame(data, columns=headers, dtype=str)
                        if len(df) > 0 and len(df.columns) > 1:
                            return df

        all_text = ''
        for page in pdf.pages:
            try:
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if text:
                    all_text += text + '\n'
            except Exception:
                pass

        if not all_text.strip():
            raise ValueError('Could not extract data from PDF. ' 'Ensure the PDF contains a BOM table or readable text.')

        lines = [l.rstrip() for l in all_text.splitlines() if l.strip()]
        if len(lines) < 2:
            raise ValueError('PDF contains insufficient rows for a BOM comparison.')

        first = lines[0]
        if '\t' in first:
            sep_pattern = r'\t'
        elif '|' in first and first.count('|') >= 2:
            sep_pattern = r'\s*\|\s*'
            lines = [l.strip('|').strip() for l in lines]
        elif ',' in first and first.count(',') >= 1:
            sep_pattern = r','
        else:
            sep_pattern = r'\s{2,}'

        split = lambda l: [c.strip() for c in re.split(sep_pattern, l)]

        headers  = split(lines[0])
        n_cols   = len(headers)
        data_rows = []
        for line in lines[1:]:
            cols = split(line)
            if all(re.match(r'^[-=]+$', c) for c in cols if c):
                continue
            # Pad / trim to match header count
            cols = (cols + [''] * n_cols)[:n_cols]
            data_rows.append(cols)

        if not data_rows:
            raise ValueError('No data rows could be parsed from the PDF text.')

        return pd.DataFrame(data_rows, columns=headers, dtype=str)


def detect_key_column(columns: list) -> Optional[str]:
    norm_to_orig = {normalize_col(c): c for c in columns}
    for alias in KEY_COLUMN_ALIASES:
        if alias in norm_to_orig:
            return norm_to_orig[alias]
    return None


def get_file_metadata(df: pd.DataFrame) -> dict:
    key_col = detect_key_column(list(df.columns))
    return {
        'row_count': len(df),
        'column_count': len(df.columns),
        'columns': list(df.columns),
        'key_column_detected': key_col or '',
    }


def normalize_value(val) -> str:
    return str(val).strip().lower()


def values_match(v1, v2) -> bool:
    return normalize_value(v1) == normalize_value(v2)


def compare_bom(master_df: pd.DataFrame, optional_df: pd.DataFrame, key_column: Optional[str] = None, master_label: str = 'Master', optional_label: str = 'Optional') -> dict:
    if key_column is None:
        key_column = (detect_key_column(list(master_df.columns)) or detect_key_column(list(optional_df.columns)))

    master_cols   = list(master_df.columns)
    optional_cols = list(optional_df.columns)

    if key_column:
        common_cols = [c for c in master_cols if c in optional_cols and c != key_column]
    else:
        common_cols = [c for c in master_cols if c in optional_cols]

    matched, partial, missing, extra = [], [], [], []

    if key_column and key_column in master_df.columns and key_column in optional_df.columns:
        master_keyed   = master_df.set_index(key_column)
        optional_keyed = optional_df.set_index(key_column)
        all_m = set(master_keyed.index.astype(str))
        all_o = set(optional_keyed.index.astype(str))

        for key in sorted(all_m & all_o):
            m_row = master_keyed.loc[key]
            o_row = optional_keyed.loc[key]
            if isinstance(m_row, pd.DataFrame): m_row = m_row.iloc[0]
            if isinstance(o_row, pd.DataFrame): o_row = o_row.iloc[0]

            col_results, all_match = {}, True
            for col in common_cols:
                mv, ov = m_row.get(col, ''), o_row.get(col, '')
                match = values_match(mv, ov)
                if not match:
                    all_match = False
                col_results[col] = {
                    'master':   str(mv),
                    'optional': str(ov),
                    'match':    bool(match),
                }

            entry = {
                'key': str(key), 'key_column': key_column,
                'column_comparison': col_results,
                'master_full':   {c: str(m_row.get(c, '')) for c in master_cols   if c != key_column},
                'optional_full': {c: str(o_row.get(c, '')) for c in optional_cols if c != key_column},
                'match_count':   sum(1 for v in col_results.values() if v['match']),
                'total_compared': len(col_results),
            }
            if all_match:
                entry['status'] = 'exact';  matched.append(entry)
            else:
                entry['status'] = 'partial'; partial.append(entry)

        for key in sorted(all_m - all_o):
            m_row = master_keyed.loc[key]
            if isinstance(m_row, pd.DataFrame): m_row = m_row.iloc[0]
            missing.append({
                'key': str(key), 'key_column': key_column, 'status': 'missing',
                'master_data': {c: str(m_row.get(c, '')) for c in master_cols if c != key_column},
            })

        for key in sorted(all_o - all_m):
            o_row = optional_keyed.loc[key]
            if isinstance(o_row, pd.DataFrame): o_row = o_row.iloc[0]
            extra.append({
                'key': str(key), 'key_column': key_column, 'status': 'extra',
                'optional_data': {c: str(o_row.get(c, '')) for c in optional_cols if c != key_column},
            })
    else:
        key_column = '#Row'
        min_rows = min(len(master_df), len(optional_df))
        for idx in range(min_rows):
            m_row, o_row = master_df.iloc[idx], optional_df.iloc[idx]
            col_results, all_match = {}, True
            for col in common_cols:
                mv, ov = m_row.get(col, ''), o_row.get(col, '')
                match = values_match(mv, ov)
                if not match: all_match = False
                col_results[col] = {'master': str(mv), 'optional': str(ov), 'match': bool(match)}
            entry = {
                'key': str(idx + 1), 'key_column': '#Row',
                'column_comparison': col_results,
                'match_count': sum(1 for v in col_results.values() if v['match']),
                'total_compared': len(col_results),
                'status': 'exact' if all_match else 'partial',
            }
            (matched if all_match else partial).append(entry)

        for idx in range(min_rows, len(master_df)):
            m_row = master_df.iloc[idx]
            missing.append({'key': str(idx + 1), 'key_column': '#Row', 'status': 'missing', 'master_data': {c: str(m_row.get(c, '')) for c in master_cols}})
        for idx in range(min_rows, len(optional_df)):
            o_row = optional_df.iloc[idx]
            extra.append({'key': str(idx + 1), 'key_column': '#Row', 'status': 'extra', 'optional_data': {c: str(o_row.get(c, '')) for c in optional_cols}})

    total = len(master_df)
    score = ((len(matched) + 0.5 * len(partial)) / total * 100) if total else 0.0

    return {
        'key_column_used': key_column or '',
        'common_columns':  common_cols,
        'match_score':     round(score, 2),
        'summary': {
            'total_master_rows':   total,
            'total_optional_rows': len(optional_df),
            'exact_matches':       len(matched),
            'partial_matches':     len(partial),
            'missing_in_optional': len(missing),
            'extra_in_optional':   len(extra),
        },
        'labels':  {'master': master_label, 'optional': optional_label},
        'details': {'matched': matched, 'partial': partial, 'missing': missing, 'extra': extra},
    }